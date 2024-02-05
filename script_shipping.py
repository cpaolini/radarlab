
import pandas as pd
import docx
import io
from docxtpl import DocxTemplate

#readingCHPAddressfileasaListstoreinContent&storedeachlineinresult
content=docx.Document(r'C:\Users\E219\Documents\certificate\area_code.docx')
shipping=docx.Document(r'C:\Users\E219\Documents\certificate\shipping.docx')
result=[p.text for p in content.paragraphs]
section = shipping.sections[0]
header = section.header
paragraphs = header.paragraphs[0]
paragraphs.text = "Date of Shipping:01/20/2024 \t No of Units: \t No of Boxes"
paragraphs.style = shipping.styles["Header"]
#shipping.add_heading('Date of Shipping', 0)
#shipping.add_heading('No of Units', 1)
#shipping.add_heading('No of Boxes', 2)
i = 0
type = input("Enter 1 for RADAR and 2 for LIDAR \n")
if type == "1":
	unit = int(input("Enter Total No of RADAR Devices \n"))
	while i<= (unit-1):
		k = input("Device Lab Number \n")
		j = input("Address Code \n")
		if(f'CHP ({j})')in result:
			ind=result.index(f'CHP ({j})')
			address_1=result[ind+2]
			address_2=result[ind+3]
			shipping.add_paragraph(f'AS24-{k}')
			shipping.add_paragraph(f'CHP ({j})')
			shipping.add_paragraph(f'Traffic Radar Coordinator')
			shipping.add_paragraph(f'{address_1}')
			shipping.add_paragraph(f'{address_2}')    	
			shipping.add_paragraph(f' ')
			shipping.save(r"C:\Users\E219\Documents\certificate\shipping.docx")
			i += 1
		else:
			print("Address Code not found \n")

elif type == "2":
	unit1 = int(input("Enter Total No of LIDAR Devices \n"))
	while i<= (unit1-1):
		l = input("Device Lab Number \n")
		m = input("Address Code \n")
		if(f'CHP ({m})')in result:
			ind=result.index(f'CHP ({m})')
			address_1=result[ind+2]
			address_2=result[ind+3]
			shipping.add_paragraph(f'ASL24-{l}')
			shipping.add_paragraph(f'CHP ({m})')
			shipping.add_paragraph(f'Traffic Radar Coordinator')
			shipping.add_paragraph(f'{address_1}')
			shipping.add_paragraph(f'{address_2}')    	
			shipping.add_paragraph(f' ')
			shipping.save(r"C:\Users\E219\Documents\certificate\shipping.docx")
			i += 1
		else:
			print("Address Code not found \n")

else:
	print("Enter 1 or 2 to make Shipping Log")
